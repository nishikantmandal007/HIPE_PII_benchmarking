import os
import json
import re
from docx import Document
from tqdm import tqdm
import nltk

# Download the sentence tokenizer model if not already present
try:
    nltk.data.find('tokenizers/punkt')
except nltk.downloader.DownloadError:
    nltk.download('punkt')

def ms_to_hms(ms):
    seconds = ms // 1000
    milliseconds = ms % 1000
    minutes = seconds // 60
    seconds = seconds % 60
    hours = minutes // 60
    minutes = minutes % 60
    return hours, minutes, seconds, milliseconds

def _process_sentences_for_srt(sentences, words_per_minute, current_time_ms):
    processed_entries = []
    i = 0
    while i < len(sentences):
        current_sentence = sentences[i].replace("\n", " ")
        words_in_current_sentence = len(current_sentence.split())
        
        # Define a threshold for "very short" sentences
        VERY_SHORT_THRESHOLD = 5 # words

        # Check if the current sentence is very short and if there's a next sentence to merge with
        if words_in_current_sentence < VERY_SHORT_THRESHOLD and i + 1 < len(sentences):
            next_sentence = sentences[i+1].replace("\n", " ")
            words_in_next_sentence = len(next_sentence.split())
            
            # Merge if the combined length is still reasonable (e.g., less than 15 words)
            # Or if both are very short
            if words_in_current_sentence + words_in_next_sentence < 15 or \
               (words_in_current_sentence < VERY_SHORT_THRESHOLD and words_in_next_sentence < VERY_SHORT_THRESHOLD):
                
                merged_sentence = current_sentence + " " + next_sentence
                words_in_merged_sentence = len(merged_sentence.split())
                duration_ms = int((words_in_merged_sentence / words_per_minute) * 60 * 1000)
                
                start_h, start_m, start_s, start_ms = ms_to_hms(current_time_ms)
                end_time_ms = current_time_ms + duration_ms
                end_h, end_m, end_s, end_ms = ms_to_hms(end_time_ms)
                
                processed_entries.append({
                    "start_time": f"{start_h:02}:{start_m:02}:{start_s:02},{start_ms:03}",
                    "end_time": f"{end_h:02}:{end_m:02}:{end_s:02},{end_ms:03}",
                    "text": merged_sentence
                })
                current_time_ms = end_time_ms
                i += 2 # Skip the next sentence as it's merged
                continue
        
        # If not merged, process as a single sentence
        duration_ms = int((words_in_current_sentence / words_per_minute) * 60 * 1000)
        
        start_h, start_m, start_s, start_ms = ms_to_hms(current_time_ms)
        end_time_ms = current_time_ms + duration_ms
        end_h, end_m, end_s, end_ms = ms_to_hms(end_time_ms)
        
        processed_entries.append({
            "start_time": f"{start_h:02}:{start_m:02}:{start_s:02},{start_ms:03}",
            "end_time": f"{end_h:02}:{end_m:02}:{end_s:02},{end_ms:03}",
            "text": current_sentence
        })
        current_time_ms = end_time_ms
        i += 1
        
    return processed_entries, current_time_ms

def generate_files(chunks, output_dir='output'):
    """
    Generate .docx, .srt, and .log files with their JSON annotations,
    organized by dataset alias and language within the specified output directory
    """
    # Create base output directories
    base_dirs = ['docs', 'docsjson', 'srt', 'srtjson', 'log', 'logjson']
    for dir_name in base_dirs:
        dir_path = os.path.join(output_dir, dir_name)
        os.makedirs(dir_path, exist_ok=True)
    
    # Get unique dataset aliases and languages
    dataset_aliases = set([chunk['dataset_alias'] for chunk in chunks])
    
    # Create subdirectories for each dataset alias and language
    for alias in dataset_aliases:
        for lang in set([c['language'] for c in chunks if c['dataset_alias'] == alias]):
            for dir_name in base_dirs:
                for split in set([c['split'] for c in chunks if c['dataset_alias'] == alias and c['language'] == lang]):
                    dir_path = os.path.join(output_dir, dir_name, alias, lang, split)
                    os.makedirs(dir_path, exist_ok=True)
    
    # Track statistics
    stats = {
        'total_chunks': len(chunks),
        'by_alias': {},
        'by_language': {},
        'by_alias_language': {}
    }
    
    # Generate files for each chunk
    for chunk in tqdm(chunks, desc="Generating files"):
        chunk_id = chunk['chunk_id']
        dataset_alias = chunk['dataset_alias']
        language = chunk['language']
        split = chunk['split']
        
        # Create a standardized ID matching the one used in entity_extraction.py
        # This ensures consistency between document IDs
        standardized_chunk_id = f"{dataset_alias}_{language}_{split}_{chunk_id}"
        
        # Use the full standardized ID for filenames to ensure consistency with entity_validation.py
        safe_chunk_id = re.sub(r'[^a-zA-Z0-9_-]', '_', standardized_chunk_id)
        
        # Update statistics
        if dataset_alias not in stats['by_alias']:
            stats['by_alias'][dataset_alias] = 0
        stats['by_alias'][dataset_alias] += 1
        
        if language not in stats['by_language']:
            stats['by_language'][language] = 0
        stats['by_language'][language] += 1
        
        alias_lang_key = f"{dataset_alias}_{language}"
        if alias_lang_key not in stats['by_alias_language']:
            stats['by_alias_language'][alias_lang_key] = 0
        stats['by_alias_language'][alias_lang_key] += 1
        
        page1_text = chunk['page1']['text']
        page2_text = chunk['page2']['text']
        
        # Create JSON annotation with person names (for docsjson and srtjson)
        page1_person_names = [entity['text'] for entity in chunk['page1']['entities']]
        page2_person_names = [entity['text'] for entity in chunk['page2']['entities']]
        
        json_annotation = {
            "page_1": page1_person_names,
            "page_2": page2_person_names
        }
        
        # Create DOCX file
        doc = Document()
        doc.add_paragraph(page1_text)
        if page2_text:  # Only add page break and second page if content exists
            doc.add_page_break()
            doc.add_paragraph(page2_text)
        doc.save(os.path.join(output_dir, "docs", dataset_alias, language, split, f"{safe_chunk_id}.docx"))
        
        # --- Create SRT file with realistic timestamps ---
        srt_content = []
        subtitle_idx = 1
        current_time_ms = 0
        words_per_minute = 180 # Average reading speed

        # Process Page 1 sentences
        sentences_page1 = nltk.sent_tokenize(page1_text)
        processed_page1_entries, current_time_ms = _process_sentences_for_srt(sentences_page1, words_per_minute, current_time_ms)
        
        for entry in processed_page1_entries:
            srt_content.append(str(subtitle_idx))
            srt_content.append(f"{entry['start_time']} --> {entry['end_time']}")
            srt_content.append(entry['text'])
            srt_content.append("") # Empty line for SRT format
            subtitle_idx += 1

        # Process Page 2 sentences (if exists)
        if page2_text:
            sentences_page2 = nltk.sent_tokenize(page2_text)
            processed_page2_entries, current_time_ms = _process_sentences_for_srt(sentences_page2, words_per_minute, current_time_ms)
            
            for entry in processed_page2_entries:
                srt_content.append(str(subtitle_idx))
                srt_content.append(f"{entry['start_time']} --> {entry['end_time']}")
                srt_content.append(entry['text'])
                srt_content.append("")
                subtitle_idx += 1

        with open(os.path.join(output_dir, "srt", dataset_alias, language, split, f"{safe_chunk_id}.srt"), 'w', encoding='utf-8') as f:
            f.write("\n".join(srt_content))

        # --- Create LOG file as plain text ---
        log_content = []
        log_content.append(f"Chunk ID: {standardized_chunk_id}")
        log_content.append(f"Dataset Alias: {dataset_alias}")
        log_content.append(f"Language: {language}")
        log_content.append(f"Split: {split}")
        log_content.append("\n--- Page 1 Text Preview ---")
        log_content.append(page1_text[:500] + "..." if len(page1_text) > 500 else page1_text) # Longer preview for logs
        if page2_text:
            log_content.append("\n--- Page 2 Text Preview ---")
            log_content.append(page2_text[:500] + "..." if len(page2_text) > 500 else page2_text)

        with open(os.path.join(output_dir, "log", dataset_alias, language, split, f"{safe_chunk_id}.log"), 'w', encoding='utf-8') as f:
            f.write("\n".join(log_content))

        # Save JSON annotations with appropriate structure for each file type
        # For docs - keep page-based structure
        with open(os.path.join(output_dir, "docsjson", dataset_alias, language, split, f"{safe_chunk_id}.json"), 'w', encoding='utf-8') as f:
            json.dump(json_annotation, f, indent=2, ensure_ascii=False)
            
        # For srt - organize entities by subtitle sections
        srt_json_annotation = {
            "metadata": {
                "chunk_id": standardized_chunk_id,
                "dataset": dataset_alias,
                "language": language,
                "split": split
            },
            "subtitles": []
        }
        
        # Group entities by subtitle
        entities_list = page1_person_names + page2_person_names
        if entities_list:
            srt_json_annotation["entities"] = entities_list
            
        with open(os.path.join(output_dir, "srtjson", dataset_alias, language, split, f"{safe_chunk_id}.json"), 'w', encoding='utf-8') as f:
            json.dump(srt_json_annotation, f, indent=2, ensure_ascii=False)
            
        # For log - create log-appropriate structure
        log_json_annotation = {
            "metadata": {
                "chunk_id": standardized_chunk_id,
                "dataset": dataset_alias,
                "language": language,
                "split": split
            },
            "entities": {}
        }
        
        # For log files, we can organize by page number as logical sections
        if page1_person_names:
            log_json_annotation["entities"]["page_1"] = page1_person_names
        if page2_person_names:
            log_json_annotation["entities"]["page_2"] = page2_person_names
            
        with open(os.path.join(output_dir, "logjson", dataset_alias, language, split, f"{safe_chunk_id}.json"), 'w', encoding='utf-8') as f:
            json.dump(log_json_annotation, f, indent=2, ensure_ascii=False)
    
    return stats